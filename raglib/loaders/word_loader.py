"""Loader for Microsoft Word (.docx) documents."""

from __future__ import annotations

import logging
import os
from typing import List

from raglib.schemas import Document

logger = logging.getLogger(__name__)


class WordLoader:
    """Loads docx files into Document objects."""

    def load(self, path: str) -> List[Document]:
        """Load a docx path into a single Document."""

        try:
            from docx import Document as DocxDocument  # type: ignore[import-not-found]
        except ImportError as exc:
            raise ImportError(
                "Word loading requires python-docx. Install with: pip install raglib[docx]"
            ) from exc

        docx_document = DocxDocument(path)
        lines = [paragraph.text.strip() for paragraph in docx_document.paragraphs if paragraph.text.strip()]

        for table in docx_document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        content = "\n".join(lines).strip()
        if not content:
            logger.warning("WordLoader found no extractable text in file=%s", path)
            return []

        document = Document(
            id=f"docx-{abs(hash(os.path.abspath(path)))}",
            content=content,
            metadata={"file": path, "loader": "word"},
            source="docx",
        )
        logger.info("WordLoader loaded 1 document from file=%s", path)
        return [
            document,
        ]
